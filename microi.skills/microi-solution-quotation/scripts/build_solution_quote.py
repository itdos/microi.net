import argparse
import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = "17365D"
BLUE = "2F75B5"
PALE_BLUE = "EAF3F9"
PALE_GRAY = "F4F6F8"
PALE_GOLD = "FFF3D6"
GOLD = "B45309"
INK = "1F2937"
MUTED = "64748B"
WHITE = "FFFFFF"
BORDER = "CBD5E1"
FONT = "Microsoft YaHei"
PAGE_WIDTH_DXA = 11909
MARGIN_DXA = 720
TABLE_INDENT_DXA = 120
TABLE_WIDTH_DXA = PAGE_WIDTH_DXA - MARGIN_DXA * 2 - TABLE_INDENT_DXA


def set_run_font(run, size=9.2, bold=False, color=INK):
    run.font.name = FONT
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        rfonts.set(qn(key), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def format_paragraph(paragraph, before=0, after=0, line=1.05, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_paragraph(doc, text, size=9.2, bold=False, color=INK, before=0, after=2, line=1.05, align=None):
    paragraph = doc.add_paragraph()
    format_paragraph(paragraph, before, after, line, align)
    set_run_font(paragraph.add_run(text), size, bold, color)
    return paragraph


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=110, bottom=70, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=TABLE_INDENT_DXA):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_section_heading(doc, text):
    return add_paragraph(doc, text, size=10.8, bold=True, color=BLUE, before=4, after=1, line=1.0)


def add_bullet(doc, label, detail):
    paragraph = doc.add_paragraph(style="List Bullet")
    format_paragraph(paragraph, before=0, after=1, line=1.03)
    set_run_font(paragraph.add_run(f"{label}："), 8.8, True, NAVY)
    set_run_font(paragraph.add_run(detail), 8.8, False, INK)


def add_metadata_table(doc, config):
    table = doc.add_table(rows=1, cols=4)
    widths = [TABLE_WIDTH_DXA // 4] * 4
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    values = [
        ("版本", config["version"]),
        ("日期", config["date"]),
        ("开发周期", config["development_cycle"]),
        ("报价有效期", f'{config["validity_days"]} 天'),
    ]
    for index, (label, value) in enumerate(values):
        cell = table.cell(0, index)
        set_cell_shading(cell, PALE_GRAY)
        paragraph = cell.paragraphs[0]
        format_paragraph(paragraph, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_run_font(paragraph.add_run(f"{label}："), 8.0, True, MUTED)
        set_run_font(paragraph.add_run(str(value)), 8.2, True, NAVY)
    set_table_geometry(table, widths)
    set_table_borders(table, "E2E8F0", "4")


def add_callout(doc, label, detail, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    set_table_borders(table, "9CC2E5", "7")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    format_paragraph(paragraph, after=0, line=1.03)
    set_run_font(paragraph.add_run(f"{label}："), 8.8, True, NAVY)
    set_run_font(paragraph.add_run(detail), 8.8, False, INK)


def add_quote_table(doc, config):
    table = doc.add_table(rows=2, cols=3)
    widths = [TABLE_WIDTH_DXA // 3] * 3
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    headers = ["市场预估价（总价）", "吾码优惠报价（总价）", "开发周期"]
    values = [config["market_price"], config["preferential_price"], config["development_cycle"]]
    for index, header in enumerate(headers):
        cell = table.cell(0, index)
        set_cell_shading(cell, NAVY)
        paragraph = cell.paragraphs[0]
        format_paragraph(paragraph, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_run_font(paragraph.add_run(header), 8.7, True, WHITE)
    for index, value in enumerate(values):
        cell = table.cell(1, index)
        set_cell_shading(cell, PALE_GOLD if index == 1 else PALE_GRAY)
        paragraph = cell.paragraphs[0]
        format_paragraph(paragraph, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_run_font(paragraph.add_run(str(value)), 12.0 if index == 1 else 10.5, True, GOLD if index == 1 else NAVY)
    set_table_geometry(table, widths)
    set_table_borders(table)


def normalize_config(raw):
    required = [
        "system_name", "project_goal", "solution_items", "technical_route",
        "deliverables", "market_price", "preferential_price",
        "development_cycle", "boundary_note",
    ]
    missing = [name for name in required if not raw.get(name)]
    if missing:
        raise ValueError("Missing required fields: " + ", ".join(missing))
    if not isinstance(raw["solution_items"], list) or not raw["solution_items"]:
        raise ValueError("solution_items must be a non-empty list")
    raw.setdefault("version", "v1.0.0")
    raw.setdefault("date", date.today().isoformat())
    raw.setdefault("validity_days", 15)
    raw.setdefault("subtitle", "AI 深度融合 V8引擎 的低代码平台，官网：https://microi.net")
    raw.setdefault("project_scope", "")
    return raw


def build_document(config, output):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = RGBColor.from_string(INK)

    title = f'Microi吾码 {config["system_name"]} 解决方案与报价'
    add_paragraph(doc, title, size=20.0, bold=True, color=NAVY, after=1, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, config["subtitle"], size=9.3, color=MUTED, after=4, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_metadata_table(doc, config)

    add_section_heading(doc, "一、项目目标")
    add_callout(doc, "目标", config["project_goal"])

    add_section_heading(doc, "二、解决方案")
    for item in config["solution_items"]:
        if isinstance(item, str):
            add_bullet(doc, "方案", item)
        else:
            add_bullet(doc, item.get("label", "方案"), item.get("detail", ""))

    add_section_heading(doc, "三、核心技术路线")
    add_callout(doc, "链路", config["technical_route"], fill=PALE_GRAY)

    add_section_heading(doc, "四、交付范围")
    add_paragraph(doc, config["deliverables"], size=8.7, after=1, line=1.03)
    if config["project_scope"]:
        add_paragraph(doc, config["project_scope"], size=8.3, color=MUTED, after=1, line=1.0)

    add_section_heading(doc, "五、项目总价与周期")
    add_quote_table(doc, config)

    add_section_heading(doc, "六、说明与边界")
    add_paragraph(doc, config["boundary_note"], size=8.1, color=MUTED, after=0, line=1.0)

    footer = section.footer.paragraphs[0]
    format_paragraph(footer, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_run_font(footer.add_run(f'Microi吾码 | {config["version"]} | {config["date"]} | https://microi.net'), 7.2, False, MUTED)

    doc.core_properties.title = title
    doc.core_properties.subject = "Microi吾码系统解决方案与报价"
    doc.core_properties.author = "Microi吾码"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main():
    parser = argparse.ArgumentParser(description="Build a compact Microi solution quotation DOCX")
    parser.add_argument("--input", required=True, help="UTF-8 JSON configuration file")
    parser.add_argument("--output", required=True, help="Output DOCX path")
    args = parser.parse_args()
    with open(args.input, "r", encoding="utf-8") as handle:
        config = normalize_config(json.load(handle))
    output = Path(args.output).resolve()
    build_document(config, output)
    print(output)


if __name__ == "__main__":
    main()
